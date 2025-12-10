import argparse
import base64
import json
import mimetypes
import os
from concurrent.futures import ThreadPoolExecutor, as_completed, wait, FIRST_COMPLETED
import shutil
from pathlib import Path
from typing import Callable, List, Tuple, Optional, Dict, Any
import re
import io
import time
import random

import fitz  # PyMuPDF
import requests
from openpyxl import Workbook
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn


def ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def is_pdf(path: Path) -> bool:
    return path.is_file() and path.suffix.lower() == ".pdf"


def iter_pdfs(input_path: Path) -> List[Path]:
    if input_path.is_file() and is_pdf(input_path):
        return [input_path]
    elif input_path.is_dir():
        return sorted([p for p in input_path.glob("**/*.pdf") if p.is_file()])
    else:
        return []


def save_embedded_images(doc: fitz.Document, out_dir: Path) -> List[Path]:
    image_paths: List[Path] = []
    for page_index in range(len(doc)):
        page = doc.load_page(page_index)
        images = page.get_images(full=True)
        for img_index, img in enumerate(images):
            xref = img[0]
            base = doc.extract_image(xref)
            img_bytes = base.get("image")
            ext = base.get("ext", "png")
            name = f"page-{page_index+1}-img-{img_index+1}.{ext}"
            img_path = out_dir / name
            with open(img_path, "wb") as f:
                f.write(img_bytes)
            image_paths.append(img_path)
    return image_paths


def render_full_pages(doc: fitz.Document, out_dir: Path, dpi: int = 200) -> List[Path]:
    page_paths: List[Path] = []
    for page_index in range(len(doc)):
        page = doc.load_page(page_index)
        zoom = dpi / 72.0
        mat = fitz.Matrix(zoom, zoom)
        pix = page.get_pixmap(matrix=mat)
        img_path = out_dir / f"page-{page_index+1}-full.png"
        pix.save(img_path)
        page_paths.append(img_path)
    return page_paths


def to_data_uri(image_path: Path) -> str:
    mime, _ = mimetypes.guess_type(str(image_path))
    if not mime:
        mime = "image/png"
    with open(image_path, "rb") as f:
        b64 = base64.b64encode(f.read()).decode("ascii")
    return f"data:{mime};base64,{b64}"


def call_doubao_extract_tables(image_path: Path, api_key: str, base_url: str, model: str) -> Tuple[str, Dict[str, int]]:
    """旧版表格提取函数，保留以兼容"""
    data_uri = to_data_uri(image_path)
    prompt = (
        "请从图片中提取所有表格，保证表格里文字内容准确无误，表格中任何列中如果有引号就直接替换成上边表格里的内容，严格输出 JSON，格式为：\n"
        "{\n  \"status\": \"ok\",\n  \"tables\": [ { \"name\": \"Table 1\", \"rows\": [[\"col1\",\"col2\"], [\"...\"]] } ]\n}\n"
        "要求：\n"
        "- 不要输出除 JSON 之外的任何文本或标记\n"
        "- 若无表格，输出 {\"status\":\"no_table\",\"tables\":[]}\n"
        "- 保留数字、小数、日期、单位等原样；合并单元格按视觉行列展开\n"
    )
    url = base_url.rstrip("/") + "/chat/completions"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
    }
    body = {
        "model": model,
        "messages": [
            {
                "role": "user",
                "content": [
                    {"type": "image_url", "image_url": {"url": data_uri}},
                    {"type": "text", "text": prompt},
                ],
            }
        ],
        # 关闭思考模式，加快响应并去掉思维链输出
        "thinking": {"type": "disabled"},
    }

    # 从环境读取可配置超时与重试次数
    read_timeout = int(os.getenv("ARK_TIMEOUT", "180"))
    retries = int(os.getenv("ARK_RETRIES", "3"))

    last_err = None
    for attempt in range(retries + 1):
        try:
            # 禁用代理（如果环境变量设置了代理但不可用，会导致连接失败）
            # 可以通过环境变量 ARK_USE_PROXY=true 来启用代理
            proxies = None
            if not os.getenv("ARK_USE_PROXY", "").lower() in ("true", "1", "yes"):
                proxies = {"http": None, "https": None}
            # 连接超时固定10秒，读取超时可配置
            resp = requests.post(url, headers=headers, json=body, timeout=(10, read_timeout), proxies=proxies)
            resp.raise_for_status()
            j = resp.json()
            break
        except (requests.exceptions.ReadTimeout, requests.exceptions.ConnectTimeout, requests.exceptions.ConnectionError) as e:
            last_err = e
            if attempt < retries:
                # 指数退避 + 抖动
                time.sleep((2 ** attempt) + random.uniform(0, 0.5))
                continue
            raise
        except requests.exceptions.RequestException:
            # 其它请求异常直接抛出
            raise
    # Try to unify text output across possible response shapes
    text = None
    usage = {}
    if isinstance(j, dict):
        if "choices" in j and j["choices"]:
            msg = j["choices"][0].get("message", {})
            content = msg.get("content", [])
            if isinstance(content, list) and content:
                parts = []
                for seg in content:
                    if isinstance(seg, dict):
                        t = seg.get("text")
                        if t:
                            parts.append(t)
                if parts:
                    text = "".join(parts)
            if not text and isinstance(msg.get("content"), str):
                text = msg["content"]
        if not text and j.get("output_text"):
            text = j.get("output_text")
    if not text:
        text = resp.text

    # 提取 usage
    if isinstance(j, dict) and isinstance(j.get("usage"), dict):
        u = j.get("usage", {})
        usage = {
            "prompt": int(u.get("prompt_tokens", 0) or 0),
            "completion": int(u.get("completion_tokens", 0) or 0),
            "total": int(u.get("total_tokens", 0) or 0),
        }
    return text, usage


def call_doubao_extract_text(image_path: Path, api_key: str, base_url: str, model: str) -> Tuple[str, Dict[str, int]]:
    """新版文字提取函数，识别所有文字并按层级结构组织"""
    data_uri = to_data_uri(image_path)
    prompt = (
        "请仔细识别图片中的所有文字内容，严格输出 JSON 格式。\n\n"
        "【适用场景】：PPT演示文稿、书本扫描、教材文档、报告文件、表格图片等各类文档。\n\n"
        "【识别要求】：\n"
        "1. 准确识别所有印刷文字，保持原文内容不变，不要遗漏\n"
        "2. 忽略以下干扰元素：水印、盖章、印章、手写批注、涂鸦、背景装饰\n"
        "3. 根据文字的格式特征判断层级结构：\n"
        "   - 字体大小、加粗、位置\n"
        "   - 中文编号：一、二、三... 或 （一）（二）... 通常是大标题\n"
        "   - 数字编号：1. 2. 3. 或 (1) (2) (3) 通常是子标题或列表\n"
        "   - 缩进和段落结构\n"
        "4. 对于表格内容，必须严格识别：\n"
        "   - 准确识别表格的所有行和列，确保列数一致\n"
        "   - 每行的单元格数量必须与表头列数完全一致\n"
        "   - 单元格内容要完整，不能遗漏或截断\n"
        "   - 严格按照视觉上的行列对齐，不能错位\n"
        "   - 如果某单元格为空，用空字符串 \"\" 表示，不能省略该单元格\n"
        "   - 表格的 rows 数组中，第一行必须是表头，后续行是数据行\n"
        "5. 按照从上到下、从左到右的阅读顺序组织内容\n\n"
        "【输出 JSON 格式】：\n"
        "{\n"
        "  \"status\": \"ok\",\n"
        "  \"content\": [\n"
        "    {\"type\": \"h1\", \"text\": \"一级大标题（如：一、xxx 或页面最大标题）\"},\n"
        "    {\"type\": \"h2\", \"text\": \"二级标题（如：(一) xxx 或 1. xxx）\"},\n"
        "    {\"type\": \"h3\", \"text\": \"三级标题（如：(1) xxx 或小节标题）\"},\n"
        "    {\"type\": \"paragraph\", \"text\": \"正文段落内容，可以很长...\"},\n"
        "    {\"type\": \"list\", \"items\": [\"列表项1\", \"列表项2\", \"列表项3\"]},\n"
        "    {\"type\": \"table\", \"rows\": [[\"表头1\",\"表头2\",\"表头3\"], [\"数据1\",\"数据2\",\"数据3\"], [\"数据4\",\"数据5\",\"数据6\"]]}\n"
        "    // 注意：table 的 rows 中，所有行的列数必须完全一致！\n"
        "  ]\n"
        "}\n\n"
        "【type 类型说明】：\n"
        "- h1: 一级大标题（页面主标题、章标题、\"一、二、三\"编号的标题）\n"
        "- h2: 二级标题（节标题、\"(一)(二)\"或\"1. 2. 3.\"编号的标题）\n"
        "- h3: 三级标题（小节标题、\"(1)(2)(3)\"编号的标题）\n"
        "- paragraph: 普通正文段落（无编号的连续文字）\n"
        "- list: 列表项（带●、•、-等符号的短条目，或连续的编号短句）\n"
        "- table: 表格数据（有明显行列结构的内容）\n"
        "  * rows 格式：[[\"列1\",\"列2\",...], [\"数据1\",\"数据2\",...], ...]\n"
        "  * 重要：所有行的列数必须完全一致，与表头列数相同\n"
        "  * 空单元格用空字符串 \"\" 表示，不能省略\n\n"
        "【重要注意事项】：\n"
        "- 只输出 JSON，不要输出任何其他文本、解释或 markdown 标记\n"
        "- 若图片无可识别文字，输出 {\"status\":\"no_text\",\"content\":[]}\n"
        "- 保留所有数字、日期、单位、标点符号原样\n"
        "- 长段落保持完整，不要拆分成多个 paragraph\n"
        "- 编号（如\"一、\"\"(1)\"）应包含在对应标题的 text 中\n"
    )
    url = base_url.rstrip("/") + "/chat/completions"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
    }
    body = {
        "model": model,
        "messages": [
            {
                "role": "user",
                "content": [
                    {"type": "image_url", "image_url": {"url": data_uri}},
                    {"type": "text", "text": prompt},
                ],
            }
        ],
        # 关闭思考模式，加快响应并去掉思维链输出
        "thinking": {"type": "disabled"},
    }

    read_timeout = int(os.getenv("ARK_TIMEOUT", "180"))
    retries = int(os.getenv("ARK_RETRIES", "3"))

    for attempt in range(retries + 1):
        try:
            # 禁用代理（如果环境变量设置了代理但不可用，会导致连接失败）
            # 可以通过环境变量 ARK_USE_PROXY=true 来启用代理
            proxies = None
            if not os.getenv("ARK_USE_PROXY", "").lower() in ("true", "1", "yes"):
                proxies = {"http": None, "https": None}
            resp = requests.post(url, headers=headers, json=body, timeout=(10, read_timeout), proxies=proxies)
            resp.raise_for_status()
            j = resp.json()
            break
        except (requests.exceptions.ReadTimeout, requests.exceptions.ConnectTimeout, requests.exceptions.ConnectionError) as e:
            if attempt < retries:
                time.sleep((2 ** attempt) + random.uniform(0, 0.5))
                continue
            raise
        except requests.exceptions.RequestException:
            raise

    text = None
    usage = {}
    if isinstance(j, dict):
        if "choices" in j and j["choices"]:
            msg = j["choices"][0].get("message", {})
            content = msg.get("content", [])
            if isinstance(content, list) and content:
                parts = []
                for seg in content:
                    if isinstance(seg, dict):
                        t = seg.get("text")
                        if t:
                            parts.append(t)
                if parts:
                    text = "".join(parts)
            if not text and isinstance(msg.get("content"), str):
                text = msg["content"]
        if not text and j.get("output_text"):
            text = j.get("output_text")
    if not text:
        text = resp.text

    if isinstance(j, dict) and isinstance(j.get("usage"), dict):
        u = j.get("usage", {})
        usage = {
            "prompt": int(u.get("prompt_tokens", 0) or 0),
            "completion": int(u.get("completion_tokens", 0) or 0),
            "total": int(u.get("total_tokens", 0) or 0),
        }
    return text, usage


def parse_model_output_to_content(text: str) -> List[Dict[str, Any]]:
    """解析模型输出，返回结构化内容列表"""
    content_list: List[Dict[str, Any]] = []
    
    # 尝试直接解析 JSON
    try:
        data = json.loads(text)
        if isinstance(data, dict) and "content" in data:
            content_list = data["content"] or []
    except Exception:
        pass
    
    # 如果上面没成功，尝试从代码块中提取 JSON
    if not content_list:
        fence_json = re.findall(r"```json\s*(.*?)```", text, flags=re.DOTALL | re.IGNORECASE)
        candidates = fence_json or re.findall(r"```\s*(.*?)```", text, flags=re.DOTALL)
        for cand in candidates:
            try:
                data = json.loads(cand)
                if isinstance(data, dict) and "content" in data:
                    content_list = data["content"] or []
                    break
            except Exception:
                continue
    
    # 如果还没成功，启发式：查找第一个 JSON 对象
    if not content_list:
        m = re.search(r"\{[\s\S]*\}", text)
        if m:
            try:
                data = json.loads(m.group(0))
                if isinstance(data, dict) and "content" in data:
                    content_list = data["content"] or []
            except Exception:
                pass
    
    # 验证和修复表格结构
    validated_list = []
    for item in content_list:
        if isinstance(item, dict) and item.get("type") == "table":
            rows = item.get("rows", [])
            if rows and len(rows) > 0:
                # 确保所有行都是列表格式
                normalized_rows = []
                for row in rows:
                    if isinstance(row, (list, tuple)):
                        normalized_rows.append([str(cell) if cell is not None else "" for cell in row])
                    else:
                        # 如果某行不是列表，跳过该行
                        continue
                
                # 统一列数：以第一行为标准
                if normalized_rows:
                    standard_cols = len(normalized_rows[0])
                    for i, row in enumerate(normalized_rows):
                        # 补齐或截断到标准列数
                        while len(row) < standard_cols:
                            row.append("")
                        normalized_rows[i] = row[:standard_cols]
                    
                    item["rows"] = normalized_rows
        validated_list.append(item)
    
    # 如果无法解析 JSON，尝试将纯文本按行分割作为段落
    if not validated_list:
        lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
        if lines:
            for line in lines:
                validated_list.append({"type": "paragraph", "text": line})
    
    return validated_list


def parse_markdown_table(md: str) -> List[List[str]]:
    lines = [ln.strip() for ln in md.splitlines() if ln.strip()]
    table_lines = [ln for ln in lines if ln.startswith("|") and ln.endswith("|")]
    if not table_lines:
        return []
    rows: List[List[str]] = []
    for ln in table_lines:
        parts = [p.strip() for p in ln.strip("|").split("|")]
        rows.append(parts)
    # Remove separator line (---)
    rows = [r for r in rows if not all(set(c) <= set("-:") for c in r)]
    return rows


def parse_model_output_to_tables(text: str) -> List[Tuple[str, List[List[str]]]]:
    tables: List[Tuple[str, List[List[str]]]] = []
    # Try JSON first
    try:
        data = json.loads(text)
        if isinstance(data, dict) and "tables" in data:
            for idx, t in enumerate(data["tables"] or []):
                name = t.get("name") or f"Table {idx+1}"
                rows = t.get("rows") or []
                if isinstance(rows, list):
                    tables.append((name, rows))
            if tables:
                return tables
    except Exception:
        # Try to extract JSON from code fences or mixed text
        # ```json ... ``` or ``` ... ```
        fence_json = re.findall(r"```json\s*(.*?)```", text, flags=re.DOTALL | re.IGNORECASE)
        candidates = fence_json or re.findall(r"```\s*(.*?)```", text, flags=re.DOTALL)
        for cand in candidates:
            try:
                data = json.loads(cand)
                if isinstance(data, dict) and "tables" in data:
                    for idx, t in enumerate(data["tables"] or []):
                        name = t.get("name") or f"Table {idx+1}"
                        rows = t.get("rows") or []
                        if isinstance(rows, list):
                            tables.append((name, rows))
                    if tables:
                        return tables
            except Exception:
                continue
        # Heuristic: grab first JSON object-like block
        m = re.search(r"\{[\s\S]*\}", text)
        if m:
            try:
                data = json.loads(m.group(0))
                if isinstance(data, dict) and "tables" in data:
                    for idx, t in enumerate(data["tables"] or []):
                        name = t.get("name") or f"Table {idx+1}"
                        rows = t.get("rows") or []
                        if isinstance(rows, list):
                            tables.append((name, rows))
                    if tables:
                        return tables
            except Exception:
                pass

    # Try Markdown tables
    md_rows = parse_markdown_table(text)
    if md_rows:
        tables.append(("Table 1", md_rows))
        return tables

    # Try CSV-like (simple)
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    if lines and "," in lines[0]:
        csv_rows: List[List[str]] = [ln.split(",") for ln in lines]
        tables.append(("Table 1", csv_rows))
        return tables

    return tables


def write_tables_to_excel(tables: List[Tuple[str, List[List[str]]]], out_path: Path) -> None:
    wb = Workbook()
    # By default, openpyxl creates a sheet named 'Sheet'; we'll replace it when writing first table
    default_sheet = wb.active
    first = True
    for idx, (name, rows) in enumerate(tables, start=1):
        if first:
            ws = default_sheet
            ws.title = name[:31] or f"Table {idx}"
            first = False
        else:
            ws = wb.create_sheet(title=(name[:31] or f"Table {idx}"))
        for row in rows:
            ws.append([str(cell) if cell is not None else "" for cell in row])
    wb.save(out_path)


def write_aggregated_excel(image_tables: List[Tuple[str, List[Tuple[str, List[List[str]]]]]], out_path: Path) -> None:
    """
    image_tables: list of (image_name, tables_per_image)
    tables_per_image: list of (table_name, rows)
    Writes all tables into a single worksheet, inserting one blank row between images.
    """
    wb = Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    
    # 如果没有任何表格数据，添加友好提示
    if not image_tables or all(not tables for _, tables in image_tables):
        ws.append(["❌ 未识别到任何表格"])
        ws.append([""])
        ws.append(["可能的原因："])
        ws.append(["1. 当前使用 embedded 模式，但PDF中的内嵌图片不包含表格"])
        ws.append(["2. 图片质量较低，模型无法识别"])
        ws.append([""])
        ws.append(["💡 建议解决方案："])
        ws.append(["• 如果是原生PDF文档，请在 .env 中设置 ARK_SOURCE=page 或 both"])
        ws.append(["• 如果是扫描版PDF，请检查图片质量或尝试提高DPI"])
        ws.append([""])
        ws.append(["详见 README.md 中的常见问题部分"])
        wb.save(out_path)
        return
    
    for image_name, tables in image_tables:
        for table_name, rows in tables:
            for row in rows:
                ws.append([str(cell) if cell is not None else "" for cell in row])
        # blank line between images
        ws.append([""])
    wb.save(out_path)


def setup_word_styles(doc: Document) -> None:
    """设置 Word 文档的样式"""
    styles = doc.styles
    
    # 设置正文样式
    try:
        normal_style = styles['Normal']
        normal_style.font.name = '微软雅黑'
        normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        normal_style.font.size = Pt(11)
    except Exception:
        pass
    
    # 设置标题样式
    heading_configs = [
        ('Heading 1', 22, True),
        ('Heading 2', 16, True),
        ('Heading 3', 14, True),
    ]
    
    for style_name, font_size, is_bold in heading_configs:
        try:
            style = styles[style_name]
            style.font.name = '微软雅黑'
            style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            style.font.size = Pt(font_size)
            style.font.bold = is_bold
        except Exception:
            pass


def write_content_to_word(all_content: List[Tuple[str, List[Dict[str, Any]]]], out_path: Path) -> None:
    """
    将所有页面的结构化内容写入 Word 文档
    all_content: list of (image_name, content_list)
    content_list: list of {"type": "h1"|"h2"|"h3"|"paragraph"|"list"|"table", ...}
    """
    doc = Document()
    setup_word_styles(doc)
    
    # 如果没有任何内容，添加友好提示
    if not all_content or all(not content for _, content in all_content):
        doc.add_paragraph("❌ 未识别到任何文字内容")
        doc.add_paragraph("")
        doc.add_paragraph("可能的原因：")
        doc.add_paragraph("1. 当前使用 embedded 模式，但PDF中的内嵌图片不包含文字")
        doc.add_paragraph("2. 图片质量较低，模型无法识别")
        doc.add_paragraph("")
        doc.add_paragraph("💡 建议解决方案：")
        doc.add_paragraph("• 如果是原生PDF文档，请在 .env 中设置 ARK_SOURCE=page 或 both")
        doc.add_paragraph("• 如果是扫描版PDF，请检查图片质量或尝试提高DPI")
        doc.save(out_path)
        return
    
    for idx, (image_name, content_list) in enumerate(all_content):
        if not content_list:
            continue
            
        for item in content_list:
            item_type = item.get("type", "paragraph")
            
            if item_type == "h1":
                text = item.get("text", "")
                if text:
                    doc.add_heading(text, level=1)
                    
            elif item_type == "h2":
                text = item.get("text", "")
                if text:
                    doc.add_heading(text, level=2)
                    
            elif item_type == "h3":
                text = item.get("text", "")
                if text:
                    doc.add_heading(text, level=3)
                    
            elif item_type == "paragraph":
                text = item.get("text", "")
                if text:
                    p = doc.add_paragraph(text)
                    
            elif item_type == "list":
                items = item.get("items", [])
                for list_item in items:
                    if list_item:
                        p = doc.add_paragraph(str(list_item), style='List Bullet')
                        
            elif item_type == "table":
                rows = item.get("rows", [])
                if rows and len(rows) > 0:
                    # 验证并统一表格结构：确保所有行的列数一致
                    # 以第一行（通常是表头）的列数为标准
                    if len(rows) > 0:
                        standard_cols = len(rows[0])
                        # 统一所有行的列数
                        normalized_rows = []
                        for row in rows:
                            normalized_row = list(row) if isinstance(row, (list, tuple)) else [str(row)]
                            # 如果列数不足，用空字符串补齐
                            while len(normalized_row) < standard_cols:
                                normalized_row.append("")
                            # 如果列数过多，截断到标准列数
                            normalized_row = normalized_row[:standard_cols]
                            normalized_rows.append(normalized_row)
                        
                        # 创建表格
                        table = doc.add_table(rows=len(normalized_rows), cols=standard_cols)
                        table.style = 'Table Grid'
                        
                        # 填充表格内容
                        for row_idx, normalized_row in enumerate(normalized_rows):
                            for col_idx, cell_text in enumerate(normalized_row):
                                cell = table.rows[row_idx].cells[col_idx]
                                cell.text = str(cell_text) if cell_text is not None else ""
                        
                        # 表格后添加空行
                        doc.add_paragraph("")
        
        # 每个图片/页面之间添加分隔（如果不是最后一个）
        if idx < len(all_content) - 1:
            doc.add_paragraph("")
    
    doc.save(out_path)


def natural_key(name: str):
    """Return a key that sorts strings in human/natural order.
    Example: img2.png < img10.png
    """
    parts = re.split(r"(\d+)", name)
    return [int(p) if p.isdigit() else p.lower() for p in parts]


def process_pdf(
    pdf_path: Path,
    output_root: Path,
    api_key: str,
    base_url: str,
    model: str,
    dpi: int = 200,
    progress_cb: Optional[Callable[[str, dict], None]] = None,
    max_workers: int = 1,
    source_mode: str = "both",
    control_getter: Optional[Callable[[], dict]] = None,
    extract_mode: str = "text",  # "text" 输出 Word，"table" 输出 Excel
) -> None:
    doc = fitz.open(str(pdf_path))
    pdf_name = pdf_path.stem
    tmp_img_dir = output_root / "tmp" / pdf_name
    
    # 根据模式选择输出目录
    if extract_mode == "table":
        output_dir = output_root / "excel"
    else:
        output_dir = output_root / "word"
    ensure_dir(tmp_img_dir)
    ensure_dir(output_dir)

    embedded_images: List[Path] = []
    full_pages: List[Path] = []
    if source_mode in ("both", "embedded"):
        embedded_images = save_embedded_images(doc, tmp_img_dir)
    if source_mode in ("both", "page"):
        full_pages = render_full_pages(doc, tmp_img_dir, dpi=dpi)

    total = len(embedded_images) + len(full_pages)
    if progress_cb:
        progress_cb(
            "start",
            {"pdf_name": pdf_name, "total": total, "embedded": len(embedded_images), "pages": len(full_pages)},
        )

    # 保持提取顺序：先按页面索引的内嵌图片，再整页渲染；不再按文件名重新排序
    images_to_process = embedded_images + full_pages
    # 构建顺序映射，确保并发完成后聚合顺序仍与迭代一致
    order_map = {p.name: idx for idx, p in enumerate(images_to_process)}
    seen = set()
    
    # 根据模式选择处理函数
    if extract_mode == "table":
        def _process_one(img_path: Path) -> Tuple[Path, Optional[str], Any]:
            try:
                raw_text, usage = call_doubao_extract_tables(img_path, api_key, base_url, model)
                tables = parse_model_output_to_tables(raw_text)
                if not tables:
                    return img_path, "no_tables", None, usage
                return img_path, None, tables, usage
            except Exception as e:
                return img_path, str(e), None, {}
        no_result_msg = "no_tables"
    else:
        def _process_one(img_path: Path) -> Tuple[Path, Optional[str], Any]:
            try:
                raw_text, usage = call_doubao_extract_text(img_path, api_key, base_url, model)
                content_list = parse_model_output_to_content(raw_text)
                if not content_list:
                    return img_path, "no_content", None, usage
                return img_path, None, content_list, usage
            except Exception as e:
                return img_path, str(e), None, {}
        no_result_msg = "no_content"

    done = 0
    image_results: List[Tuple[str, Any]] = []
    usage_totals = {"prompt": 0, "completion": 0, "total": 0}
    
    def _wait_if_paused():
        if not control_getter:
            return False
        while True:
            ctrl = control_getter() or {}
            if ctrl.get("stop"):
                return True
            if not ctrl.get("paused"):
                return False
            time.sleep(0.5)

    if max_workers and max_workers > 1:
        with ThreadPoolExecutor(max_workers=max_workers) as ex:
            imgs_iter = (p for p in images_to_process if p not in seen)
            running = set()
            # 初始提交
            while len(running) < max_workers:
                if _wait_if_paused():
                    break
                nxt = next(imgs_iter, None)
                if not nxt:
                    break
                running.add(ex.submit(_process_one, nxt))
            # 处理循环
            while running:
                done_set, _ = wait(running, timeout=0.5, return_when=FIRST_COMPLETED)
                for fut in done_set:
                    running.remove(fut)
                    img_path, err, result_data, usage_delta = fut.result()
                    seen.add(img_path)
                    done += 1
                    if usage_delta:
                        usage_totals["prompt"] += usage_delta.get("prompt", 0) or 0
                        usage_totals["completion"] += usage_delta.get("completion", 0) or 0
                        usage_totals["total"] += usage_delta.get("total", 0) or 0
                    if progress_cb:
                        progress_cb("step", {"pdf_name": pdf_name, "done": done, "total": total, "image": img_path.name, "error": err, "usage": usage_delta})
                    if err:
                        print(f"[WARN] 处理 {img_path.name} 失败: {err}")
                    else:
                        print(f"[OK] {pdf_name}: {img_path.name}")
                        if result_data:
                            image_results.append((img_path.name, result_data))
                    # 尝试提交下一张
                    if _wait_if_paused():
                        running.clear()
                        break
                    nxt = next(imgs_iter, None)
                    if nxt:
                        running.add(ex.submit(_process_one, nxt))
    else:
        for img_path in images_to_process:
            if img_path in seen:
                continue
            if _wait_if_paused():
                break
            seen.add(img_path)
            img_path, err, result_data, usage_delta = _process_one(img_path)
            done += 1
            if usage_delta:
                usage_totals["prompt"] += usage_delta.get("prompt", 0) or 0
                usage_totals["completion"] += usage_delta.get("completion", 0) or 0
                usage_totals["total"] += usage_delta.get("total", 0) or 0
            if progress_cb:
                progress_cb("step", {"pdf_name": pdf_name, "done": done, "total": total, "image": img_path.name, "error": err, "usage": usage_delta})
            if err:
                print(f"[WARN] 处理 {img_path.name} 失败: {err}")
            else:
                print(f"[OK] {pdf_name}: {img_path.name}")
                if result_data:
                    image_results.append((img_path.name, result_data))

    if progress_cb:
        progress_cb("finish", {"pdf_name": pdf_name, "done": done, "total": total, "usage": usage_totals})

    # 根据模式输出不同格式
    image_results_sorted = sorted(image_results, key=lambda x: order_map.get(x[0], 10**9))
    if extract_mode == "table":
        excel_out = output_dir / f"{pdf_name}.xlsx"
        write_aggregated_excel(image_results_sorted, excel_out)
    else:
        word_out = output_dir / f"{pdf_name}.docx"
        write_content_to_word(image_results_sorted, word_out)

    # 清理临时图片目录
    try:
        shutil.rmtree(tmp_img_dir)
    except Exception:
        pass


def main():
    parser = argparse.ArgumentParser(description="从 PDF/图片中识别文字，自动分级标题，输出 Word 文档")
    parser.add_argument("--input", required=True, help="PDF 文件或包含 PDF 的目录路径")
    parser.add_argument("--output", default="output", help="输出根目录，默认 output")
    parser.add_argument("--model", default="doubao-seed-1-6-vision-250815", help="模型名")
    parser.add_argument("--dpi", type=int, default=200, help="整页渲染 DPI，默认 200")
    parser.add_argument("--source", choices=["both", "embedded", "page"], default=os.getenv("ARK_SOURCE", "both"), help="图片来源：both/embedded/page")
    args = parser.parse_args()

    api_key = os.getenv("ARK_API_KEY")
    base_url = os.getenv("ARK_BASE_URL")
    if not api_key or not base_url:
        raise RuntimeError(
            "缺少 Ark API 配置信息，请设置环境变量 ARK_API_KEY 与 ARK_BASE_URL"
        )

    input_path = Path(args.input)
    output_root = Path(args.output)
    ensure_dir(output_root)

    pdfs = iter_pdfs(input_path)
    if not pdfs:
        raise RuntimeError(f"未找到 PDF：{input_path}")

    print(f"发现 {len(pdfs)} 个 PDF，开始处理……")
    workers = int(os.getenv("ARK_WORKERS", "1"))
    for pdf in pdfs:
        print(f"处理: {pdf}")
        process_pdf(
            pdf,
            output_root,
            api_key,
            base_url.rstrip("/"),
            args.model,
            dpi=args.dpi,
            progress_cb=None,
            max_workers=max(1, workers),
            source_mode=args.source,
        )
    print("✅ 处理完成！Word 文档已生成到 output/word/ 目录")


def process_images(
    image_paths: List[Path],
    batch_name: str,
    output_root: Path,
    api_key: str,
    base_url: str,
    model: str,
    progress_cb: Optional[Callable[[str, dict], None]] = None,
    max_workers: int = 1,
    control_getter: Optional[Callable[[], dict]] = None,
    extract_mode: str = "text",  # "text" 输出 Word，"table" 输出 Excel
) -> None:
    # 根据模式选择输出目录
    if extract_mode == "table":
        output_dir = output_root / "excel"
    else:
        output_dir = output_root / "word"
    ensure_dir(output_dir)

    # 直接使用上传的图片路径进行处理，不再复制或保存
    # 多图批处理按自然文件名顺序排序，确保 img2 在 img10 之前
    images_to_process: List[Path] = sorted(list(image_paths), key=lambda p: natural_key(p.name))
    order_map = {p.name: idx for idx, p in enumerate(images_to_process)}

    total = len(images_to_process)
    if progress_cb:
        progress_cb("start", {"pdf_name": batch_name, "total": total, "embedded": total, "pages": 0})

    # 根据模式选择处理函数
    if extract_mode == "table":
        def _process_one(img_path: Path) -> Tuple[Path, Optional[str], Any, Dict[str, int]]:
            try:
                raw_text, usage = call_doubao_extract_tables(img_path, api_key, base_url, model)
                tables = parse_model_output_to_tables(raw_text)
                if not tables:
                    return img_path, "no_tables", None, usage
                return img_path, None, tables, usage
            except Exception as e:
                return img_path, str(e), None, {}
    else:
        def _process_one(img_path: Path) -> Tuple[Path, Optional[str], Any, Dict[str, int]]:
            try:
                raw_text, usage = call_doubao_extract_text(img_path, api_key, base_url, model)
                content_list = parse_model_output_to_content(raw_text)
                if not content_list:
                    return img_path, "no_content", None, usage
                return img_path, None, content_list, usage
            except Exception as e:
                return img_path, str(e), None, {}

    done = 0
    image_results: List[Tuple[str, Any]] = []
    seen = set()
    usage_totals = {"prompt": 0, "completion": 0, "total": 0}
    
    def _wait_if_paused():
        if not control_getter:
            return False
        while True:
            ctrl = control_getter() or {}
            if ctrl.get("stop"):
                return True
            if not ctrl.get("paused"):
                return False
            time.sleep(0.5)

    if max_workers and max_workers > 1:
        with ThreadPoolExecutor(max_workers=max_workers) as ex:
            imgs_iter = (p for p in images_to_process if p not in seen)
            running = set()
            while len(running) < max_workers:
                if _wait_if_paused():
                    break
                nxt = next(imgs_iter, None)
                if not nxt:
                    break
                running.add(ex.submit(_process_one, nxt))
            while running:
                done_set, _ = wait(running, timeout=0.5, return_when=FIRST_COMPLETED)
                for fut in done_set:
                    running.remove(fut)
                    img_path, err, result_data, usage_delta = fut.result()
                    seen.add(img_path)
                    done += 1
                    if usage_delta:
                        usage_totals["prompt"] += usage_delta.get("prompt", 0) or 0
                        usage_totals["completion"] += usage_delta.get("completion", 0) or 0
                        usage_totals["total"] += usage_delta.get("total", 0) or 0
                    if progress_cb:
                        progress_cb("step", {"pdf_name": batch_name, "done": done, "total": total, "image": img_path.name, "error": err, "usage": usage_delta})
                    if not err and result_data:
                        image_results.append((img_path.name, result_data))
                    if _wait_if_paused():
                        running.clear()
                        break
                    nxt = next(imgs_iter, None)
                    if nxt:
                        running.add(ex.submit(_process_one, nxt))
    else:
        for img_path in images_to_process:
            if img_path in seen:
                continue
            if _wait_if_paused():
                break
            seen.add(img_path)
            img_path, err, result_data, usage_delta = _process_one(img_path)
            done += 1
            if usage_delta:
                usage_totals["prompt"] += usage_delta.get("prompt", 0) or 0
                usage_totals["completion"] += usage_delta.get("completion", 0) or 0
                usage_totals["total"] += usage_delta.get("total", 0) or 0
            if progress_cb:
                progress_cb("step", {"pdf_name": batch_name, "done": done, "total": total, "image": img_path.name, "error": err, "usage": usage_delta})
            if not err and result_data:
                image_results.append((img_path.name, result_data))

    if progress_cb:
        progress_cb("finish", {"pdf_name": batch_name, "done": done, "total": total, "usage": usage_totals})

    # 根据模式输出不同格式
    image_results_sorted = sorted(image_results, key=lambda x: order_map.get(x[0], 10**9))
    if extract_mode == "table":
        excel_out = output_dir / f"{batch_name}.xlsx"
        write_aggregated_excel(image_results_sorted, excel_out)
    else:
        word_out = output_dir / f"{batch_name}.docx"
        write_content_to_word(image_results_sorted, word_out)


if __name__ == "__main__":
    main()
